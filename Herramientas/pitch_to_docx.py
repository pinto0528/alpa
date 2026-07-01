#!/usr/bin/env python3
"""Convert pitch-v2.md to pitch-v2.docx — clean, minimal formatting."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os, re

MD_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "scratch")
OUT_DIR = MD_DIR

MD_PATH = os.path.join(MD_DIR, "pitch-v2.md")
OUT_PATH = os.path.join(OUT_DIR, "pitch-v2.docx")

def parse_inline(text):
    parts = []
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*)')
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            parts.append((text[last:m.start()], False, False))
        if m.group(2):
            parts.append((m.group(2), True, False))
        elif m.group(3):
            parts.append((m.group(3), False, True))
        last = m.end()
    if last < len(text):
        parts.append((text[last:], False, False))
    return parts

def main():
    with open(MD_PATH, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    for sec in doc.sections:
        sec.page_width = Inches(8.5)
        sec.page_height = Inches(11)
        sec.top_margin = Inches(0.8)
        sec.bottom_margin = Inches(0.8)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(4)

    for line in lines:
        raw = line.rstrip('\n').rstrip('\r')

        if not raw.strip():
            continue

        if raw.strip() == '---' or raw.strip() == '___':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('_' * 60)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run.font.size = Pt(8)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            continue

        h_match = re.match(r'^(#{1,5})\s+(.+)$', raw)
        if h_match:
            level = len(h_match.group(1))
            text = h_match.group(2)
            p = doc.add_paragraph()
            if level <= 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sizes = {1: 24, 2: 20, 3: 16, 4: 14, 5: 12}
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(sizes.get(level, 12))
            if level == 3:
                run.font.color.rgb = RGBColor(0x91, 0x1B, 0x1E)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            continue

        p = doc.add_paragraph()
        for seg_text, b, italic in parse_inline(raw):
            run = p.add_run(seg_text)
            if b:
                run.bold = True
            if italic:
                run.italic = True

    doc.save(OUT_PATH)
    print(f"[OK] Pitch generado: {OUT_PATH}")

if __name__ == '__main__':
    main()
