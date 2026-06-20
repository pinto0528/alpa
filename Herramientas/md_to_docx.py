#!/usr/bin/env python3
"""Convert entregable_2.md to entregable_2.docx matching entregable_1 format exactly."""

import re, os, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

MD_PATH = os.path.join("Entregables_concurso", "Entregable_2", "entregable_2.md")
OUT_PATH = os.path.join("Entregables_concurso", "Entregable_2", "entregable_2.docx")

def resolve_image(md_dir, rel_path):
    full = os.path.normpath(os.path.join(md_dir, rel_path))
    return full if os.path.isfile(full) else None

def parse_inline(text):
    """Parse **bold**, *italic* into segments (text, bold, italic)."""
    parts = []
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*)')
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            parts.append((text[last:m.start()], False, False))
        if m.group(2):
            parts.append((m.group(2), True, False))
        elif m.group(3):
            parts.append((m.group(3), False, True))
        last = m.end()
    if last < len(text):
        parts.append((text[last:], False, False))
    return parts

def is_table_row(line):
    return line.startswith('|') and line.endswith('|') and '|' in line[1:-1] if line.strip() else False

def is_table_separator(line):
    s = line.strip()
    return bool(re.match(r'^\|[\s\-:|+]+\|$', s)) if s else False

def add_formatted_run(paragraph, text, bold=False, italic=False, size=None):
    """Add a run with optional formatting. Returns the run."""
    run = paragraph.add_run(text)
    if size:
        run.font.size = Pt(size)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return run

def write_table(doc, rows_data):
    """Write a table matching E1 style: Table Grid, 10pt font, bold only from inline markers."""
    if not rows_data:
        return
    # Filter out empty rows (all cells are empty/whitespace) like E1 does
    clean = []
    for row in rows_data:
        if any(c.strip() for c in row):
            clean.append(row)
    if not clean:
        return
    ncols = max(len(r) for r in clean)
    table = doc.add_table(rows=len(clean), cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(clean):
        for j, cell_text in enumerate(row_data):
            if j >= ncols:
                break
            cell = table.cell(i, j)
            cell.text = ''
            p = cell.paragraphs[0]
            for seg_text, b, italic in parse_inline(cell_text):
                run = p.add_run(seg_text)
                run.font.size = Pt(10)
                if b:
                    run.bold = True
                if italic:
                    run.italic = True
    doc.add_paragraph()  # spacer

def process_md():
    base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    md_path = os.path.join(base, MD_PATH)
    out_path = os.path.join(base, OUT_PATH)

    if not os.path.isfile(md_path):
        print(f"[ERR] No se encuentra: {md_path}")
        sys.exit(1)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # ---- Page setup: Letter + 1 inch margins (match E1) ----
    for sec in doc.sections:
        sec.page_width = Inches(8.5)
        sec.page_height = Inches(11)
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)

    # ---- Normal style: Calibri, inherit size, no spacing (match E1) ----
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = None  # inherit from theme
    # Remove space_after
    pPr = normal.element.get_or_add_pPr()
    for spacing in pPr.findall(qn('w:spacing')):
        pPr.remove(spacing)

    md_dir = os.path.dirname(os.path.abspath(md_path))
    i = 0
    table_buffer = []
    past_cover_hr = False  # first --- passed
    past_body_hr = False    # second --- passed

    def flush_table():
        nonlocal table_buffer
        if table_buffer:
            write_table(doc, table_buffer)
            table_buffer = []

    while i < len(lines):
        line = lines[i].rstrip('\n').rstrip('\r')

        # Empty lines → just an empty paragraph for spacing (like E1)
        if not line.strip():
            flush_table()
            # Add empty paragraph with spacing similar to E1
            p = doc.add_paragraph()
            i += 1
            continue

        # ---- Horizontal rule ----
        if line.strip() == '---':
            flush_table()
            if not past_cover_hr:
                past_cover_hr = True
            else:
                past_body_hr = True
            # `---` does not produce visible output; blank lines around it handle spacing
            i += 1
            continue

        # ---- Blockquote ----
        if line.startswith('> '):
            flush_table()
            text = line[2:]
            p = doc.add_paragraph()
            for seg_text, b, italic in parse_inline(text):
                run = p.add_run(seg_text)
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                if b:
                    run.bold = True
                if italic:
                    run.italic = True
            pPr = p._p.get_or_add_pPr()
            ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="567"/>')
            pPr.append(ind)
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                '  <w:left w:val="single" w:sz="12" w:space="8" w:color="1A3A5C"/>'
                '</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue

        # ---- Image ----
        img_match = re.match(r'!\[(.*?)\]\((.+?)\)', line.strip())
        if img_match:
            flush_table()
            rel_path = img_match.group(2)
            abs_path = resolve_image(md_dir, rel_path)
            if abs_path:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(abs_path, width=Inches(5.5))
            i += 1
            continue

        # ---- Table ----
        if is_table_row(line):
            cells = [c.strip() for c in line.strip('|').split('|')]
            if is_table_separator(line):
                i += 1
                continue
            table_buffer.append(cells)
            i += 1
            # Check next line
            if i < len(lines):
                nxt = lines[i].rstrip('\n').rstrip('\r')
                if not is_table_row(nxt) and not is_table_separator(nxt):
                    flush_table()
            continue

        flush_table()

        # ---- Unordered list ----
        stripped = line.strip()
        ul_match = re.match(r'^[-*+]\s+(.+)$', stripped)
        if ul_match:
            flush_table()
            text = ul_match.group(1)
            p = doc.add_paragraph()
            try:
                p.style = doc.styles['List Bullet']
            except KeyError:
                pass
            for seg_text, b, italic in parse_inline(text):
                run = p.add_run(seg_text)
                if b:
                    run.bold = True
                if italic:
                    run.italic = True
            i += 1
            continue

        # ---- Ordered list ----
        ol_match = re.match(r'^(\d+)[.)]\s+(.+)$', stripped)
        if ol_match:
            flush_table()
            text = ol_match.group(2)
            p = doc.add_paragraph()
            try:
                p.style = doc.styles['List Number']
            except KeyError:
                pass
            for seg_text, b, italic in parse_inline(text):
                run = p.add_run(seg_text)
                if b:
                    run.bold = True
                if italic:
                    run.italic = True
            i += 1
            continue

        # ---- Heading detection ----
        stripped = line.strip()
        heading_match = re.match(r'^(#{1,5})\s+(.+)$', stripped)

        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)

            # Skip cover masthead lines (# EmprendeU, ## Concurso...) before first ---
            if not past_cover_hr:
                i += 1
                continue

            p = doc.add_paragraph()

            # Determine size & formatting matching E1's pattern
            if level == 3 and not past_body_hr:
                # ### Entregable N° N → cover title: sz=44 Bold, #1B3A5C
                hsize = 44
                bold = True
                color = '1B3A5C'
                spBefore, spAfter = 240, 80
            elif level == 4 and not past_body_hr:
                # #### Módulo N... → cover subtitle: sz=26 NOT Bold, #555555
                hsize = 26
                bold = False
                color = '555555'
                spBefore, spAfter = None, 80
            elif level == 2:
                # ## Section → sz=30 Bold, #2E568B
                hsize = 30
                bold = True
                color = '2E568B'
                spBefore, spAfter = 240, 80
            elif level == 3:
                # ### Subsection → sz=26 Bold, #3A6A9C
                hsize = 26
                bold = True
                color = '3A6A9C'
                spBefore, spAfter = 160, 80
            else:
                hsize = 22
                bold = True
                color = '3A6A9C'
                spBefore, spAfter = None, None

            for seg_text, b, italic in parse_inline(text):
                run = p.add_run(seg_text)
                run.font.size = Pt(hsize // 2)
                run.font.color.rgb = RGBColor(*[int(color[i:i+2], 16) for i in (0, 2, 4)])
                if bold or b:
                    run.bold = True
                if italic:
                    run.italic = True

            if spBefore is not None or spAfter is not None:
                pPr = p._p.get_or_add_pPr()
                attrs = ''
                if spBefore is not None:
                    attrs += f' w:before="{spBefore}"'
                if spAfter is not None:
                    attrs += f' w:after="{spAfter}"'
                spacing = parse_xml(f'<w:spacing {nsdecls("w")}{attrs}/>')
                pPr.append(spacing)

            i += 1
            continue

        # ---- Regular paragraph (one per line, like E1) ----
        p = doc.add_paragraph()
        for seg_text, b, italic in parse_inline(line):
            run = p.add_run(seg_text)
            if b:
                run.bold = True
            if italic:
                run.italic = True
        i += 1

    flush_table()
    doc.save(out_path)
    print(f"[OK] Documento generado: {out_path}")

if __name__ == '__main__':
    process_md()
